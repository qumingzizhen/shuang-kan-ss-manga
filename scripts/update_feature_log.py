from __future__ import annotations

import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT = PROJECT_ROOT / "docs" / "项目功能变更记录.docx"


ENTRIES = [
    {
        "date": "2026-06-11",
        "area": "总体架构",
        "feature": "建立 newwork 新架构工作区",
        "details": "创建 web-first 的 monorepo 结构：apps/web、services/api、workers/download、packages/domain、infra、docs。明确网页端先行，未来 App 与网页端共用同一个后端。",
        "status": "已完成",
        "verification": "项目结构已落地，README 与架构文档已记录。",
    },
    {
        "date": "2026-06-11",
        "area": "本地环境",
        "feature": "D 盘项目本地缓存与开发脚本",
        "details": "新增 scripts/dev-env.ps1，把 npm、Cargo、Rustup、临时目录等缓存引导到 <project-root> 下，避免把依赖和临时文件散到 C 盘。",
        "status": "已完成",
        "verification": "scripts/check.ps1 多次通过；环境输出显示缓存目录位于项目内。",
    },
    {
        "date": "2026-06-11",
        "area": "前端",
        "feature": "Next.js 网页控制台",
        "details": "创建 Next.js/React/TypeScript 前端控制台，包含任务创建、任务列表、统计指标、事件日志、搜索/直链/补缺三个输入模式。",
        "status": "已完成",
        "verification": "npm --prefix .\\apps\\web run build 通过。",
    },
    {
        "date": "2026-06-11",
        "area": "前端检索",
        "feature": "多 tag 输入解析",
        "details": "tag 输入支持中文逗号、英文逗号、中文分号、英文分号、换行和空格作为分隔符；同时保留 female:big breasts 这类命名空间 tag 的完整语义。",
        "status": "已完成",
        "verification": "TypeScript 构建通过，逻辑已接入搜索任务创建。",
    },
    {
        "date": "2026-06-11",
        "area": "后端",
        "feature": "Rust Axum API 骨架",
        "details": "创建 Rust/Axum API，提供健康检查、任务列表、搜索任务、图库直链任务、补缺任务、单任务查询、任务事件等接口。",
        "status": "已完成",
        "verification": "cargo check --workspace 通过。",
    },
    {
        "date": "2026-06-11",
        "area": "领域模型",
        "feature": "任务领域模型与请求契约",
        "details": "在 packages/domain 中定义 Task、TaskKind、TaskStatus、TaskPayload、TaskProgress 以及创建任务请求结构，为前后端和 worker 后续协作打基础。",
        "status": "已完成",
        "verification": "Rust workspace check 通过。",
    },
    {
        "date": "2026-06-11",
        "area": "数据层",
        "feature": "任务仓储边界与 PostgreSQL 实现",
        "details": "后端支持默认内存仓储，并实现 PostgreSQL 仓储作为可选 feature。已添加 tasks 表迁移。",
        "status": "部分完成",
        "verification": "默认内存模式 check 通过；PostgreSQL 代码路径受当前 Windows GNU 链接器限制，尚未完成 API 运行验证。",
    },
    {
        "date": "2026-06-11",
        "area": "数据库",
        "feature": "项目本地 PostgreSQL 17.10",
        "details": "下载并解压 PostgreSQL Windows zip 到 .tools，数据目录放在 .data，新增 scripts/postgres.ps1 管理启动、状态、连接串和本地数据库初始化。Docker Desktop 暂缓安装。",
        "status": "已完成",
        "verification": "psql 已连接 manga 数据库，manga 用户可用，tasks 表迁移已执行。",
    },
    {
        "date": "2026-06-11",
        "area": "事件流",
        "feature": "内存任务事件总线与 SSE",
        "details": "新增 TaskEvent 模型、内存发布器历史缓存和广播通道，API 暴露 /v1/tasks/events，前端通过 EventSource 订阅任务入队与更新事件。",
        "status": "已完成",
        "verification": "scripts/check.ps1 通过，包括 Rust format、Rust workspace check、Next.js build。",
    },
    {
        "date": "2026-06-11",
        "area": "文档",
        "feature": "项目功能变更记录 Word 文档",
        "details": "创建本 Word 文档，用于持续记录项目每次新增或修改的功能、影响范围、状态和验证结果。",
        "status": "已完成",
        "verification": "已生成 docs/项目功能变更记录.docx。",
    },
    {
        "date": "2026-06-11",
        "area": "任务控制",
        "feature": "任务生命周期更新与取消",
        "details": "新增 PATCH /v1/tasks/{id} 用于更新任务状态、进度和标题；新增 POST /v1/tasks/{id}/cancel 用于取消任务。前端任务表增加取消按钮，任务更新会发布 task_updated 事件并通过 SSE 同步到控制台。",
        "status": "已完成",
        "verification": "scripts/check.ps1 已通过，包括 Rust workspace check 与 Next.js build。",
    },
    {
        "date": "2026-06-11",
        "area": "扩展架构",
        "feature": "源站适配器注册与能力校验",
        "details": "新增 SourceAdapterDescriptor 与 SourceCapability，API 暴露 GET /v1/sources。搜索、直链和补缺任务创建时会校验所选来源是否支持对应能力；前端任务表单从接口加载来源列表，不再写死来源。新增 docs/source-adapters.md 记录后续适配器扩展规则。",
        "status": "已完成",
        "verification": "scripts/check.ps1 已通过，包括 Rust workspace check 与 Next.js build。",
    },
    {
        "date": "2026-06-12",
        "area": "扩展架构",
        "feature": "共享源站适配器 SDK 与 worker 分发器",
        "details": "新增 packages/source-adapter 共享包，定义 SourceAdapter trait、适配器注册表、图库摘要、页面描述、下载产物和补缺计划等扩展契约。API 和下载 worker 共同使用该注册表；worker 新增 TaskDispatcher，将任务 payload 分发到对应来源适配器，为后续接入更多网站、插件进程或 WASM 适配器预留空间。",
        "status": "已完成",
        "verification": "scripts/check.ps1 已通过，包括新增 source-adapter 包、API 和下载 worker。",
    },
    {
        "date": "2026-06-12",
        "area": "队列架构",
        "feature": "共享任务队列契约与内存队列实现",
        "details": "新增 packages/task-queue 共享包，定义 TaskQueue trait、TaskQueueMessage、QueueError 和 InMemoryTaskQueue。API 创建任务时会在写入仓储后入队，再发布任务事件；下载 worker 也初始化队列契约，为后续替换成 NATS JetStream 跨进程队列预留空间。新增 docs/task-queue.md 记录消息契约和未来 NATS subject 规划。",
        "status": "已完成",
        "verification": "scripts/check.ps1 已通过，包括新增 task-queue 包、API 入队逻辑和下载 worker。",
    },
    {
        "date": "2026-06-12",
        "area": "Worker 架构",
        "feature": "下载 worker 运行时与任务 reporter 边界",
        "details": "新增 WorkerRuntime 和 TaskReporter。worker 现在围绕 TaskQueueMessage 执行“接收消息、上报开始、分发任务、上报完成或失败、确认消息”的标准流程。当前 TracingTaskReporter 写入日志，后续可替换为 API 或数据库 reporter 来持久更新任务状态和进度，而不影响来源适配器。",
        "status": "已完成",
        "verification": "scripts/check.ps1 已通过，包括 worker runtime、task reporter 与队列 ack/retry 扩展。",
    },
    {
        "date": "2026-06-12",
        "area": "任务生命周期",
        "feature": "任务事件分类与领域生命周期助手",
        "details": "将任务事件从单一 task_updated 扩展为 task_started、task_progressed、task_completed、task_failed、task_canceled 等明确事件；在 packages/domain 中加入 TaskProgress 校验以及 Task 状态、进度、重命名、取消等生命周期方法。API 更新和取消任务时会按语义发布对应事件，前端事件流同步监听这些事件。",
        "status": "已完成",
        "verification": "scripts/check.ps1 已通过，包括 Rust workspace check 与 Next.js build；功能日志脚本通过 Python 语法检查并重新生成 Word 文档。",
    },
    {
        "date": "2026-06-12",
        "area": "任务执行",
        "feature": "共享任务运行时与 API 本地内置 worker",
        "details": "新增 packages/task-runtime，将 TaskDispatcher、WorkerRuntime、TaskReporter 与 TracingTaskReporter 从下载 worker 中抽成共享包。API 默认启动本地内置 worker，能在内存队列模式下消费任务、回写 running/completed/failed 状态，并通过生命周期事件同步给前端；后续接入 NATS 时可切换到独立下载 worker。",
        "status": "已完成",
        "verification": "cargo check --workspace 已通过；worker 中旧 dispatcher/runtime 文件已移除，独立 worker 和 API 共用同一套运行时包。",
    },
    {
        "date": "2026-06-12",
        "area": "源站适配",
        "feature": "Fangliding 旧 Python 下载器桥接 adapter",
        "details": "新增 scripts/fangliding_bridge.py，并让内置 Fangliding adapter 通过该桥调用旧 ex_fangliding_downloader.py。当前桥接覆盖搜索、图库元数据读取和已有目录补缺计划；完整逐页下载与细粒度进度回传仍留作下一阶段。",
        "status": "部分完成",
        "verification": "Python 语法检查通过；使用旧 .venv 对已有下载目录执行 retry-plan 只读测试，成功返回 JSON；Rust workspace check 已通过。",
    },
    {
        "date": "2026-06-12",
        "area": "下载链路",
        "feature": "图库直链任务执行整本下载",
        "details": "扩展 SourceAdapter 契约，新增 download_gallery 和 GalleryDownloadReport。TaskDispatcher 在处理图库直链任务时会要求来源同时具备 gallery 与 download 能力，并调用 Fangliding bridge 的 download-gallery 命令执行旧下载器整本下载。完成后会把 total、done、failed 写入 TaskProgress，前端可看到真实完成统计。默认输出目录位于 <project-root>\\.data\\downloads。",
        "status": "已完成",
        "verification": "Python 语法检查通过；Rust workspace check 通过；download-gallery 非法 URL 校验测试按预期失败；已有目录 retry-plan 只读测试继续成功返回 JSON。",
    },
    {
        "date": "2026-06-12",
        "area": "任务结果",
        "feature": "结构化任务输出与搜索结果一键下载",
        "details": "新增 TaskOutput 与 TaskSearchResult，搜索任务会保存 search_results，图库直链任务会保存 gallery_download，补缺任务会保存 retry_plan。API 本地 worker 在任务完成时把结构化结果写回任务快照；PostgreSQL 迁移新增 output JSONB 字段。前端任务列表会展示搜索到的图库，并可从单条结果直接创建图库下载任务。",
        "status": "已完成",
        "verification": "scripts/check.ps1 已通过，包括 Rust format、默认 Rust workspace check 和 Next.js build；API PostgreSQL feature 额外检查受当前 Windows GNU dlltool CreateProcess 环境限制未完成。",
    },
    {
        "date": "2026-06-12",
        "area": "本地运行",
        "feature": "开发期 Node.js API shim",
        "details": "新增 services/dev-api/server.mjs，提供与网页控制台兼容的临时 API：健康检查、来源列表、任务列表、SSE 事件流、搜索任务、图库直链下载任务、补缺任务、任务更新和取消。shim 通过 scripts/fangliding_bridge.py 调用旧 Fangliding 下载器桥，作为 Rust API 在当前 Windows 链接器问题解决前的本地可运行入口；正式后端路线仍保留 Rust Axum API。",
        "status": "已完成",
        "verification": "node --check 已通过；scripts/check.ps1 已通过；临时启动 shim 后，/health 与 /v1/sources 可访问，使用已有下载目录创建 retry-folder 任务可完成并返回 retry_plan 输出。Rust cargo run 仍受 GNU dlltool/缺少 assembler 与本机无 VS Build Tools 限制。",
    },
    {
        "date": "2026-06-12",
        "area": "本地运行",
        "feature": "一键开发启动与 shim 任务落盘",
        "details": "新增 scripts/dev.ps1 和 npm run dev，用一个终端启动开发期 API shim 与 Next.js 网页控制台；Ctrl+C 时会停止 API shim。开发期 API shim 新增任务快照持久化，completed 任务和结构化 output 会保存到 <project-root>\\.data\\dev-api\\tasks.json，重启后仍可在任务列表中看到。若重启时存在 queued/running/paused 任务，会标记为 failed 并提示需要重新创建任务。",
        "status": "已完成",
        "verification": "node --check、PowerShell 语法检查与 scripts/check.ps1 均通过；烟测创建 retry-folder 任务后重启 shim，任务仍可从 .data/dev-api/tasks.json 读取并保留 retry_plan 输出；测试数据已清理。",
    },
    {
        "date": "2026-06-25",
        "area": "本地运行",
        "feature": "开发启动脚本端口复用与强制重启",
        "details": "scripts/dev.ps1 不再因为 8080 或 3000 已被本项目占用就直接失败。脚本会用 /health 识别已运行的开发 API，并在网页控制台也已运行时直接输出可打开地址；需要加载最新代码时可显式使用 .\\scripts\\dev.ps1 -Fresh 重启对应端口。端口占用检测改为 TCP 探测，减少受限环境下 Get-NetTCPConnection 不可用导致的误判；API 子进程启动改为构造干净环境，规避 Windows 环境变量 PATH/Path 重复造成的 Start-Process 异常。",
        "status": "已完成",
        "verification": "在 8080 与 3000 均已占用的状态下执行 .\\scripts\\dev.ps1，脚本成功复用现有服务并输出 http://127.0.0.1:3000；scripts/check.ps1 已通过。",
    },
    {
        "date": "2026-06-12",
        "area": "前端任务体验",
        "feature": "搜索结果多选批量下载与任务详情侧栏",
        "details": "任务列表中的 search_results 输出现在支持勾选、全选、清空和批量创建图库下载任务；单条结果仍可直接创建下载任务。新增任务详情侧栏，可查看任务类型、状态、进度、错误/完成信息、搜索结果全集、下载输出目录、补缺页码、payload JSON 和 output JSON，并支持复制输出路径或 JSON。",
        "status": "已完成",
        "verification": "scripts/check.ps1 已通过，包括 Dev API shim 语法检查、Rust format、Rust workspace check 与 Next.js build/type check。",
    },
    {
        "date": "2026-06-12",
        "area": "文件库",
        "feature": "只读本地文件库雏形",
        "details": "开发期 API shim 新增 GET /v1/library，只读扫描 <project-root>\\.data\\downloads 与旧下载目录 <legacy-download-root>，可识别 metadata.json、failed_pages.jsonl、图片数量、页数、tag、图库 URL、目录大小和更新时间。网页端新增侧边栏文件库视图，展示本地漫画目录统计、失败记录、路径、tag，并支持复制目录或元数据路径。",
        "status": "已完成",
        "verification": "node --check 通过；scripts/check.ps1 已通过；临时启动 dev API shim 后请求 /v1/library，成功读到旧下载目录中的 1 个漫画条目，image_count/page_count 均为 752，failed_count 为 0。",
    },
    {
        "date": "2026-06-12",
        "area": "文件库",
        "feature": "文件库详情侧栏与图片预览流",
        "details": "开发期 API shim 新增 GET /v1/library/{id} 和 GET /v1/library/{id}/pages/{filename}。详情接口返回单本漫画的 metadata 摘要、页面文件列表和失败记录；图片接口只读流式返回扫描根目录下直接漫画子目录中的图片文件，并校验 id、文件名和扩展名。网页端文件库表格新增查看详情按钮，侧栏按需加载单本详情，展示本地目录、来源 URL、tag、失败记录、metadata JSON 和前 24 页图片预览。",
        "status": "已完成",
        "verification": "node --check 与 scripts/check.ps1 已通过；临时启动 dev API shim 后请求 /v1/library/{id} 成功返回 752 个页面条目，请求第一张图片返回 200、Content-Type 为 image/webp、大小为 96992 字节。",
    },
    {
        "date": "2026-06-12",
        "area": "文件库",
        "feature": "文件库检索、筛选与排序",
        "details": "开发期 API shim 的 GET /v1/library 新增 q、tag、completeness、failed_only、sort 查询参数，可按标题/路径/tag 文本、独立 tag、完整/缺页状态、失败记录和排序条件返回库存。网页端文件库新增关键词、tag、完整度、仅失败和排序控件，统计卡片与表格会随当前筛选结果即时更新，并提供清空筛选操作。该参数契约后续可下沉到 PostgreSQL 或 Meilisearch。",
        "status": "已完成",
        "verification": "node --check 与 scripts/check.ps1 已通过；临时启动 dev API shim 后请求 /v1/library?q=skill&tag=artist&completeness=complete&sort=title_asc 返回 1 个漫画条目，请求 /v1/library?failed_only=true 返回 0 个条目。",
    },
    {
        "date": "2026-06-13",
        "area": "文件库导出",
        "feature": "CBZ 导出",
        "details": "新增 scripts/library_export.py，使用 Python 标准库 zipfile 将本地漫画目录导出为 CBZ，并保留 metadata.json 与 failed_pages.jsonl。开发期 API shim 新增 POST /v1/library/{id}/exports/cbz，输出目录固定在 <project-root>\\.data\\exports\\cbz。网页端文件库详情侧栏新增导出 CBZ 按钮，导出完成后展示文件路径、页数、体积并支持复制路径。",
        "status": "已完成",
        "verification": "node --check、Python 语法检查与 scripts/check.ps1 已通过；临时启动 dev API shim 后对旧下载目录执行 CBZ 导出，生成的 archive 可被 zipfile 打开，图片条目数为 752，并包含 metadata.json 与 failed_pages.jsonl。",
    },
    {
        "date": "2026-06-13",
        "area": "文件库导出",
        "feature": "PDF 导出",
        "details": "新增 scripts/library_pdf_export.py，使用 Pillow 读取本地图片并通过 ReportLab 生成逐页 PDF。开发期 API shim 新增 POST /v1/library/{id}/exports/pdf，输出目录固定在 <project-root>\\.data\\exports\\pdf。网页端文件库详情侧栏新增导出 PDF 按钮，并与 CBZ 一样展示导出路径、页数、体积和复制路径操作。",
        "status": "已完成",
        "verification": "node --check、Python 语法检查与 scripts/check.ps1 已通过；临时启动 dev API shim 后对旧下载目录执行 PDF 导出，生成的 PDF 文件头校验为 %PDF，导出页数为 752。",
    },
    {
        "date": "2026-06-13",
        "area": "文件库修复",
        "feature": "文件库详情创建补缺任务",
        "details": "网页端文件库详情侧栏新增补缺区，可基于当前漫画本地目录直接创建 missing-only 补缺任务。创建成功后会复用既有 retry-folder 任务流，自动切回任务控制台并打开新任务详情，避免重复实现补缺逻辑。",
        "status": "已完成",
        "verification": "scripts/check.ps1 已通过；临时启动 dev API shim 后以 UTF-8 JSON 调用 POST /v1/tasks/retry-folder，任务完成并返回 retry_plan，样例目录待补页数为 0。",
    },
    {
        "date": "2026-06-13",
        "area": "文件库导出",
        "feature": "导出历史记录",
        "details": "开发期 API shim 新增 GET /v1/library/{id}/exports，并在每次 CBZ/PDF 导出成功后向 <project-root>\\.data\\dev-api\\library-exports.jsonl 追加 manifest 记录。网页端文件库详情会读取并展示导出历史，刷新页面后仍可看到历史输出路径、格式、页数、体积和文件存在状态。",
        "status": "已完成",
        "verification": "node --check 与 scripts/check.ps1 已通过；临时启动 dev API shim 后执行 CBZ 导出，再请求 /v1/library/{id}/exports，可返回包含该导出的历史记录。",
    },
    {
        "date": "2026-06-13",
        "area": "文件库导出",
        "feature": "导出文件网页下载",
        "details": "开发期 API shim 新增 GET /v1/library/{id}/exports/{exportId}/file，用附件方式流式返回 manifest 中记录的 CBZ/PDF 文件。接口会校验导出记录属于当前漫画，且文件仍位于允许的项目本地导出目录。网页端导出历史为存在的文件显示下载按钮，避免直接暴露任意本地路径。",
        "status": "已完成",
        "verification": "node --check 与 scripts/check.ps1 已通过；临时启动 dev API shim 后请求导出文件下载接口，返回 200、正确 Content-Type 和附件文件名响应头。",
    },
    {
        "date": "2026-06-13",
        "area": "文件库导出",
        "feature": "导出历史刷新与时间显示",
        "details": "网页端文件库详情中的导出历史增加创建时间显示和手动刷新按钮，可在不关闭详情侧栏的情况下重新读取 manifest 中的导出记录。",
        "status": "已完成",
        "verification": "前端 TypeScript 构建通过；功能复用既有 GET /v1/library/{id}/exports 接口。",
    },
    {
        "date": "2026-06-13",
        "area": "文件库导出",
        "feature": "导出下载链接复制",
        "details": "网页端导出历史为仍存在的 CBZ/PDF 导出记录增加复制下载链接按钮，便于调试、复用或手动打开受控下载接口。",
        "status": "已完成",
        "verification": "前端 TypeScript 构建通过。",
    },
    {
        "date": "2026-06-13",
        "area": "文件库预览",
        "feature": "预览图打开原图",
        "details": "网页端文件库详情中的页面预览缩略图现在可点击打开原图流，仍复用既有 GET /v1/library/{id}/pages/{filename} 安全图片接口。",
        "status": "已完成",
        "verification": "前端 TypeScript 构建通过。",
    },
    {
        "date": "2026-06-13",
        "area": "文件库预览",
        "feature": "页面预览显示更多",
        "details": "网页端文件库详情的页面预览从固定前 24 页扩展为可分批显示更多页面，并提供收起按钮，便于在不引入复杂虚拟滚动前检查长篇漫画页面。",
        "status": "已完成",
        "verification": "前端 TypeScript 构建通过。",
    },
    {
        "date": "2026-06-13",
        "area": "文件库预览",
        "feature": "页面清单分页加载",
        "details": "开发期 API shim 新增 GET /v1/library/{id}/pages?offset=0&limit=24，按批返回页面元数据。详情接口只返回首批页面和分页总数，网页端显示更多时再请求下一批，避免大图库打开详情时一次性传输全部页面清单。",
        "status": "已完成",
        "verification": "node --check、前端 TypeScript 构建与 scripts/check.ps1 已通过；临时启动 dev API shim 后分页请求返回 total=752、首批 24 条、next_offset=24。",
    },
    {
        "date": "2026-06-13",
        "area": "文件库检索",
        "feature": "tag 点击筛选",
        "details": "网页端文件库列表和详情中的 tag 现在可点击，点击后会自动填入文件库 tag 筛选条件并回到文件库列表，减少手动复制 tag 的操作。",
        "status": "已完成",
        "verification": "前端 TypeScript 构建通过。",
    },
    {
        "date": "2026-06-13",
        "area": "文件库修复",
        "feature": "列表一键创建补缺任务",
        "details": "网页端文件库列表的操作列新增创建补缺任务按钮，可直接基于当前漫画目录创建 missing-only retry-folder 任务，无需先进入详情侧栏。",
        "status": "已完成",
        "verification": "前端 TypeScript 构建通过。",
    },
    {
        "date": "2026-06-13",
        "area": "前端任务体验",
        "feature": "任务列表筛选",
        "details": "网页端任务控制台新增任务筛选面板，可按关键词、任务类型和任务状态过滤任务列表。关键词会匹配任务标题、ID、进度信息、payload 和 output，便于在历史任务变多后快速定位搜索、直链下载或补缺任务。",
        "status": "已完成",
        "verification": "前端 TypeScript 构建通过。",
    },
    {
        "date": "2026-06-13",
        "area": "前端任务体验",
        "feature": "任务重跑",
        "details": "网页端任务列表和任务详情侧栏新增重跑入口，会从旧任务保存的 payload 还原搜索、直链下载或补缺任务的创建请求，并重新调用既有创建接口生成新任务。该功能避免重复手填参数，也保持与正式 API 的任务校验逻辑一致。",
        "status": "已完成",
        "verification": "前端 TypeScript 构建通过。",
    },
    {
        "date": "2026-06-13",
        "area": "任务接口",
        "feature": "任务列表查询参数",
        "details": "Rust API 与开发期 API shim 的 GET /v1/tasks 新增 q、kind、status 查询参数，可按任务标题、ID、类型、状态、进度、payload 和 output 文本过滤任务列表。前端 SDK 同步增加 TaskListParams，为后续 PostgreSQL 下推过滤、任务历史分页和公开平台任务检索预留稳定契约。",
        "status": "已完成",
        "verification": "node --check、Rust workspace check 与前端 TypeScript 构建通过。",
    },
    {
        "date": "2026-06-13",
        "area": "前端任务体验",
        "feature": "任务统计快捷筛选",
        "details": "网页端任务控制台的统计卡片改为可点击按钮。点击全部任务、排队、运行或失败卡片会直接切换任务列表的状态筛选，并通过高亮边框显示当前筛选状态，减少在下拉框中重复选择的操作。",
        "status": "已完成",
        "verification": "前端 TypeScript 构建通过。",
    },
    {
        "date": "2026-06-13",
        "area": "文件库书架",
        "feature": "本地书架与阅读状态",
        "details": "开发期 API shim 新增 PATCH /v1/library/{id}/shelf，并将收藏、阅读状态、备注和更新时间持久化到 <project-root>\\.data\\dev-api\\library-shelf.json。文件库列表和详情会显示书架信息，支持收藏切换、阅读状态修改、备注保存、仅收藏筛选和阅读状态筛选，为后续账号系统中的用户书架表预留数据形态。",
        "status": "已完成",
        "verification": "node --check、前端 TypeScript 构建、接口烟测与 scripts/check.ps1 通过。",
    },
    {
        "date": "2026-06-14",
        "area": "前端交互与视觉",
        "feature": "侧边栏收回逻辑与二次元风格调整",
        "details": "任务详情和文件库详情侧边栏新增更明确的“收回”按钮，并增加点击侧边栏外部区域自动收回的遮罩交互。整体网页控制台从偏默认后台样式调整为更柔和的漫画/二次元风格，包括纸面感背景、粉青配色、漫画感描边、彩色按钮和更醒目的抽屉层级。",
        "status": "已完成",
        "verification": "前端 TypeScript 构建和 scripts/check.ps1 通过。",
    },
    {
        "date": "2026-06-14",
        "area": "文件库阅读体验",
        "feature": "阅读进度与继续阅读",
        "details": "本地书架数据新增 last_page 和 last_read_at，记录读到第几页和最后阅读时间。文件库列表显示阅读进度并提供继续阅读按钮；详情侧栏显示进度条、继续阅读、清除进度，并在页面预览中提供“读到这里”标记。文件库统计卡片也升级为快捷筛选，可一键查看全部、失败、在读和收藏条目；侧边栏支持 Escape 收回。",
        "status": "已完成",
        "verification": "node --check、前端 TypeScript 构建、书架接口烟测与 scripts/check.ps1 通过。",
    },
    {
        "date": "2026-06-14",
        "area": "文件库阅读体验",
        "feature": "内置漫画阅读器",
        "details": "网页端文件库新增全屏内置漫画阅读器，列表继续阅读和详情页缩略图点击都会进入同一阅读界面。阅读器支持上一页、下一页、适应宽度、适应高度、原始大小、键盘翻页和 Escape 收起；每次打开或翻页都会复用 PATCH /v1/library/{id}/shelf 写回 last_page 与阅读状态，使列表、详情侧栏和阅读器进度保持一致。",
        "status": "已完成",
        "verification": "npm --prefix .\\apps\\web run build 已通过；后续已纳入 scripts/check.ps1 综合检查。",
    },
    {
        "date": "2026-06-14",
        "area": "文件库阅读体验",
        "feature": "阅读器页码跳转与预加载",
        "details": "内置漫画阅读器新增页码输入跳转、当前页附近页码快捷按钮、Home/End 快捷跳到首页和末页，并在阅读当前页时自动预取前后两页元数据与图片资源。该能力仍复用 GET /v1/library/{id}/pages 分页接口，不改变后端契约，后续可自然扩展为阅读器缩略图条、章节目录或移动端手势翻页。",
        "status": "已完成",
        "verification": "前端 TypeScript 构建通过；后续已纳入 scripts/check.ps1 综合检查。",
    },
    {
        "date": "2026-06-14",
        "area": "文件库阅读体验",
        "feature": "阅读器缩略图条与偏好记忆",
        "details": "内置漫画阅读器新增当前页前后三页的缩略图胶片条，当前页高亮，点击缩略图可直接跳转对应页。阅读器的适应宽度、适应高度、原始大小模式会写入浏览器 localStorage，下次进入阅读器时自动恢复偏好。该功能仅增强网页端体验，不改变后端接口和书架数据契约。",
        "status": "已完成",
        "verification": "前端 TypeScript 构建通过；后续已纳入 scripts/check.ps1 综合检查。",
    },
    {
        "date": "2026-06-14",
        "area": "文件库阅读体验",
        "feature": "最近阅读书架与启动提示优化",
        "details": "文件库新增最近阅读统计卡片、最近阅读快捷筛选、last_read_desc 排序、recent_only 接口参数，以及按 last_read_at 排列的继续阅读书架。用户可以在文件库顶部直接回到最近读过的漫画，也可以只筛选有阅读历史的条目。开发 API shim 同步支持 recent_only 与 last_read_desc，后续迁移到 PostgreSQL 或搜索索引时可沿用同一接口契约。scripts/check.ps1 结束时新增提示，明确该脚本只做检查，并给出 .\\scripts\\dev.ps1 启动命令。",
        "status": "已完成",
        "verification": "node --check、前端 TypeScript 构建与 scripts/check.ps1 均已通过；Word 文档重新生成并完成渲染检查。",
    },
    {
        "date": "2026-06-14",
        "area": "文件库视觉与浏览",
        "feature": "本地漫画封面与卡片视图",
        "details": "开发期 API shim 在文件库条目中新增 cover_filename 与 cover_url，从本地漫画目录按自然页序选取第一张图片作为封面，并继续通过现有 /v1/library/{id}/pages/{filename} 安全图片流返回，不暴露任意本地路径。网页端文件库新增表格/卡片视图切换，默认使用封面卡片视图；表格行、最近阅读书架和卡片都能显示封面，并复用同一套继续阅读、详情、复制目录、补缺和收藏操作。",
        "status": "已完成",
        "verification": "node --check、前端 TypeScript 构建与 scripts/check.ps1 均已通过；Word 文档重新生成并完成渲染检查。",
    },
    {
        "date": "2026-06-15",
        "area": "文件库书架管理",
        "feature": "批量选择与批量书架操作",
        "details": "网页端文件库新增多选状态与批量操作栏，可全选当前筛选结果、清空选择，并对选中漫画批量收藏、取消收藏或修改阅读状态。表格视图和封面卡片视图都提供选择框，选中条目会有明确视觉状态。批量更新复用现有 PATCH /v1/library/{id}/shelf 接口，先保证本地可用与逻辑一致，后续可升级为真正的批量 API。",
        "status": "已完成",
        "verification": "前端 TypeScript 构建与 scripts/check.ps1 均已通过；Word 文档重新生成并完成渲染检查。",
    },
    {
        "date": "2026-06-15",
        "area": "源站适配器安全边界",
        "feature": "登录保护源站与 ExHentai 类站点接入原则",
        "details": "补充 docs/source-adapters.md，明确需要账号、cookie、年龄门槛、区域授权或其他访问边界的源站只能使用用户/操作者合法授权的本地 cookie 或凭据；不得绕过登录、付费墙、CAPTCHA、封禁或显式访问控制。ExHentai 类站点的安全路线是本地配置授权 cookie、低并发读取元数据和页面 URL，并先以只读 adapter 方式接入，确认边界后再考虑下载能力。",
        "status": "已完成",
        "verification": "文档已更新；scripts/check.ps1 已通过；Word 文档重新生成并完成渲染检查。",
    },
    {
        "date": "2026-06-15",
        "area": "源站适配器",
        "feature": "18comic.vip 源站爬取适配器",
        "details": "新增 scripts/18comic_bridge.py，支持 18comic.vip 的公开搜索、图库元数据解析、保守串行下载和已有目录补缺计划。Rust SourceAdapterRegistry 注册新的 18comic 内置来源；开发期 Node API shim 改为多来源注册表，可按 source_id 分流到 Fangliding 或 18comic 桥接脚本。文件库 source_id 现在优先读取 metadata.json，避免多源下载目录被误标为默认来源。",
        "status": "已完成初版",
        "verification": "18comic 桥接脚本通过 py_compile 与离线 self-test；dev-api server.mjs 通过 node --check；scripts/check.ps1 已通过并纳入 18comic 桥接脚本自测。后续仍需在用户可正常访问 18comic.vip 的网络环境中做真实搜索和小批量下载 smoke test。",
    },
    {
        "date": "2026-06-15",
        "area": "源站适配器架构",
        "feature": "通用 Python 桥接适配器与共享爬虫内核",
        "details": "将 Fangliding 与 18comic 的 Rust 内置 adapter 重构为同一个 PythonBridgeAdapter，内置源站变为 descriptor、脚本路径、Python 运行时和能力开关的配置项；开发期 Node API shim 同步改为 sourceAdapters 单表配置，避免 sources 与 bridgeScripts 两份清单漂移。新增 scripts/source_bridge_core.py，集中 HTTP 请求、重试退避、cookie/header 文件读取、HTML 基础抽取、图片校验、安全文件名、原子 JSON 写入等公共能力，站点脚本只保留 URL 规则与页面解析策略。",
        "status": "已完成",
        "verification": "scripts/check.ps1 已通过，覆盖 JSON 配置、dev-api 语法、source_bridge_core.py 与两个桥接脚本的 py_compile、18comic 离线 self-test、Rust format、Rust workspace check 和 Next.js 构建；/v1/sources 冒烟测试确认前端只看到统一 source descriptor。",
    },
    {
        "date": "2026-06-15",
        "area": "源站适配器架构",
        "feature": "共享源站注册表配置",
        "details": "新增 config/source-adapters.json，集中声明默认源站、source descriptor、能力列表、桥接脚本路径、脚本环境变量、Python 运行时环境变量和 page-level 命令开关。Rust source-adapter 包默认内嵌读取该配置，并支持 SOURCE_ADAPTER_CONFIG 指向外部注册表；开发期 Node API shim 同样读取该配置派生 /v1/sources 与内部桥接分发。新增 scripts/check_source_adapters.py 校验注册表，避免新增源站时出现重复 ID、未知能力、缺失脚本或错误桥接类型。",
        "status": "已完成",
        "verification": "scripts/check.ps1 已通过，新增覆盖 config/source-adapters.json JSON 解析与 scripts/check_source_adapters.py 注册表校验；Rust workspace check、dev-api 语法检查、Next.js 构建均通过；临时启动 dev-api 后 /v1/sources 返回两个统一 descriptor 且不暴露 bridge 内部字段。",
    },
    {
        "date": "2026-06-15",
        "area": "代码质量与可扩展性",
        "feature": "代码结构与可复用性审查重构",
        "details": "完成一次项目级代码结构审查，确认 Rust 后端核心已基本形成领域模型、源站 Adapter/Registry、任务队列、运行时、仓储与事件发布等分层；同时识别网页端 Dashboard 和开发期 API shim 单文件职责过重。本轮新增 docs/代码结构与可复用性审查.md，将 Dashboard 中的 tag 解析、状态标签、任务重跑、文件库排序和阅读进度计算抽到 apps/web/src/lib/dashboard-model.ts，并将 dev-api 的源站注册表加载、校验和 bridge 物化抽到 services/dev-api/source-registry.mjs。",
        "status": "已完成",
        "verification": "node --check services/dev-api/server.mjs 与 services/dev-api/source-registry.mjs 通过；npm --prefix .\\apps\\web run build 通过；scripts/check.ps1 已加入新模块语法检查。",
    },
    {
        "date": "2026-06-15",
        "area": "文件库检索",
        "feature": "热门 tag 统计与快捷筛选",
        "details": "开发期 API shim 新增 GET /v1/library/tags，可按当前文件库筛选条件统计热门 tag，返回 tag、命中漫画数、图片总数和失败页总数。前端 SDK 新增 LibraryTagStat 与 listLibraryTags，并在网页文件库中加入热门 tag 面板，点击 tag 会复用现有 tag 筛选逻辑。该契约后续可迁移到 PostgreSQL 聚合查询或 Meilisearch facet。",
        "status": "已完成",
        "verification": "node --check services/dev-api/server.mjs 通过；npm --prefix .\\apps\\web run build 通过；临时启动 dev-api 后请求 /v1/library/tags?limit=5 成功返回 tag 统计 JSON。",
    },
    {
        "date": "2026-06-16",
        "area": "文件库导出",
        "feature": "多选批量 CBZ/PDF 导出",
        "details": "网页端文件库的批量操作栏新增批量 CBZ 与批量 PDF 按钮。用户可以先在表格或封面卡片中多选漫画，再顺序调用现有单本 CBZ/PDF 导出接口生成多个导出文件；每个成功结果会写回前端导出历史缓存，后续打开详情侧栏仍能看到对应记录。当前实现优先复用已有端点，后续可下沉为真正的批量导出 API 或后台任务队列。",
        "status": "已完成",
        "verification": "npm --prefix .\\apps\\web run build 通过；node --check services/dev-api/server.mjs 通过；功能复用既有 CBZ/PDF 导出接口，未新增后端路由。",
    },
    {
        "date": "2026-06-16",
        "area": "前端交互与源站任务体验",
        "feature": "侧栏收回动画与默认多源爬取",
        "details": "任务详情和文件库详情侧栏的收回逻辑从立即卸载改为 closing 状态加延迟卸载，遮罩和侧栏都会执行淡出/滑出动画，并保留 Escape、收回按钮和点击外部区域三种关闭入口。任务输入的来源选择默认改为“全部源站一起爬取”，提交时按所有启用 source descriptor 展开为多个任务；用户只有在下拉框手动选择具体源站时才只爬取一个站。",
        "status": "已完成",
        "verification": "npm --prefix .\\apps\\web run build 通过；node --check services/dev-api/server.mjs 通过；前端仍通过 GET /v1/sources 获取启用源站，不硬编码站点数量。",
    },
    {
        "date": "2026-06-16",
        "area": "源站适配与检索稳定性",
        "feature": "18comic 搜索词归一化与 403 候选入口容错",
        "details": "修复 18comic 搜索中直接使用 female:big breasts、language:chinese 等命名空间 tag 造成源站查询语法不匹配的问题。18comic bridge 现在会在源站适配层把命名空间 tag 转换为普通关键词，例如 female:big breasts 转为 big breasts，并在单个公开搜索入口返回 401、403 或 429 时继续尝试同源其它公开搜索入口；只有所有候选入口都被拒绝时才返回明确错误。该逻辑保留在源站 bridge 内部，前端和任务模型仍保持多源中立。",
        "status": "已完成",
        "verification": "python scripts/18comic_bridge.py self-test 通过；python -m py_compile scripts/18comic_bridge.py scripts/source_bridge_core.py scripts/check_source_adapters.py 通过；python scripts/check_source_adapters.py 通过。",
    },
    {
        "date": "2026-06-16",
        "area": "检索体验与任务编排",
        "feature": "多源搜索结果合并输出",
        "details": "搜索模式下选择全部源站时，前端不再为每个源站分别创建独立搜索任务，而是向开发期 API shim 提交一个带 source_ids 的多源搜索任务。后端会逐个调用对应 source bridge，按 source_id 与 gallery_url 去重合并结果，并在同一份 search_results 输出中保留每条结果的来源，后续批量下载仍可按结果自身 source_id 分流到正确源站。若部分源站失败，任务仍合并显示其它成功源的结果，并在结果区显示轻量提示。",
        "status": "已完成",
        "verification": "node --check services/dev-api/server.mjs 通过；npm --prefix .\\apps\\web run build 通过；scripts/check.ps1 通过。",
    },
    {
        "date": "2026-06-25",
        "area": "检索体验",
        "feature": "搜索结果源站缩略图",
        "details": "搜索结果数据契约新增 thumbnail_url 字段。source_bridge_core 提供通用的链接邻近图片提取辅助函数，可从 img、懒加载属性、srcset、CSS url() 和裸图片 URL 中找到靠近图库链接的缩略图；18comic 与 e-hentai 搜索解析器接入该能力。缩略图候选会过滤下载箭头、icon、sprite、placeholder 以及 E-Hentai 的 ehgt.org/g/t.png、ehgt.org/g/td.png 等站点 UI 图标，避免把按钮图标当成漫画封面。开发期 API shim 合并多源搜索结果时会保留并清洗 thumbnail_url，并新增 GET /v1/search-thumbnails 缩略图代理缓存接口：前端只访问本地 API，后端校验 source_id、远程 URL、来源域名和私网地址，首次拉取后缓存到 .data\\thumbnail-cache，后续直接读取本地缓存。网页端任务列表和任务详情侧栏都会显示固定比例缩略图，旧任务中的错误缩略图也会退回稳定占位。该字段属于通用 search_results 输出，后续新增源站只需要在适配器返回同名字段即可复用。",
        "status": "已完成",
        "verification": "python bridge py_compile、18comic/e-hentai self-test、node --check services/dev-api/server.mjs、npm --prefix .\\apps\\web run build 与 scripts/check.ps1 均已通过。",
    },
    {
        "date": "2026-06-17",
        "area": "架构与数据库设计",
        "feature": "项目架构与数据库设计文档",
        "details": "新增 docs/项目架构与数据库设计.md 与 docs/项目架构与数据库设计.docx，系统化记录 Web/App 共用后端、源站适配器、在线阅读链路、PostgreSQL 主库、Redis 缓存、对象存储、tag 翻译、多词条 tag 检索和公开化治理等设计。该文档用于后续功能扩展时对齐整体架构，避免每加一个源站或功能就出现孤立组件。",
        "status": "已完成",
        "verification": "使用 scripts/generate_design_doc.py 从 Markdown 生成同名 Word 文档。",
    },
    {
        "date": "2026-06-17",
        "area": "在线阅读与源站抽象",
        "feature": "搜索结果直接进入在线阅读器",
        "details": "开发期 API shim 新增 /v1/reader/sessions 远程阅读会话接口，以及 /v1/reader/sessions/{id}/pages 与 /v1/reader/sessions/{id}/pages/{index} 页列表和单页图片接口。源站能力扩展为 page_list、page_image、online_read；18comic 通过已有 list-pages 与 download-page 桥接命令接入按页阅读，图片按需缓存到 .data/page-cache。网页端搜索结果新增阅读按钮，支持从搜索结果直接打开内置阅读器，不需要先下载整本漫画。",
        "status": "已完成",
        "verification": "node --check services/dev-api/server.mjs 通过；python -m py_compile 通过；python scripts/18comic_bridge.py self-test 通过；python scripts/check_source_adapters.py 通过；npm --prefix .\\apps\\web run build 通过；scripts/check.ps1 通过。",
    },
    {
        "date": "2026-06-17",
        "area": "在线阅读体验",
        "feature": "远程阅读会话持久化与直链在线阅读",
        "details": "开发期 API shim 新增 reader-sessions.json 仓储，远程阅读会话会保存 source、gallery_url、title、tag、页列表和更新时间，服务重启后仍可恢复页列表；同时开放 GET /v1/reader/sessions 与 GET /v1/reader/sessions/{id}，为后续“继续上次在线阅读”功能预留接口。网页端直链模式新增“在线阅读”按钮，会在当前选择范围内自动挑选支持 online_read 的源站并打开内置阅读器，下载任务与立即阅读入口分离，减少用户必须先下载再看的不适感。",
        "status": "已完成",
        "verification": "node --check services/dev-api/server.mjs 通过；python -m py_compile 通过；python scripts/check_source_adapters.py 通过；python scripts/18comic_bridge.py self-test 通过；npm --prefix .\\apps\\web run build 通过；scripts/check.ps1 通过。",
    },
    {
        "date": "2026-06-17",
        "area": "在线阅读体验",
        "feature": "在线阅读进度保存与继续阅读面板",
        "details": "开发期 API shim 新增 PATCH /v1/reader/sessions/{id}/progress，用于保存远程阅读会话的 last_page 与 last_read_at；reader-sessions.json 快照现在会持久化阅读进度，重新创建同一源站同一图库的会话时会保留已有进度。网页端启动时加载最近远程阅读会话，并在任务视图右侧新增“最近在线阅读”面板，展示页码进度、百分比进度条和“继续”按钮；从搜索结果、直链入口或翻页进入阅读页时才更新进度，普通图片预取不会污染阅读位置。",
        "status": "已完成",
        "verification": "node --check services/dev-api/server.mjs 通过；npm --prefix .\\apps\\web run build 通过；隔离 dev-api 烟测确认 PATCH progress 可把伪会话从第 1 页更新到第 2 页并返回 last_read_at；scripts/check.ps1 通过；Word 文档重新生成并完成渲染检查。",
    },
    {
        "date": "2026-06-19",
        "area": "在线阅读体验",
        "feature": "远程阅读器临近页预载状态",
        "details": "网页端远程在线阅读器的临近页预取从纯后台行为升级为可见状态：打开或翻到某一页时，前端会拉取当前页前后若干页的 page descriptor，并用浏览器 Image 对象预热对应图片缓存；阅读器底部新增预载状态胶囊，可显示预载中、已预载和预载失败。预载失败只记录日志并显示为预载状态，不再弹出全局错误，也不会调用 PATCH progress，因此不会污染真实阅读页码。",
        "status": "已完成",
        "verification": "python -m py_compile scripts/update_feature_log.py scripts/generate_design_doc.py scripts/check_source_adapters.py 通过；node --check services/dev-api/server.mjs 与 services/dev-api/source-registry.mjs 通过；npm --prefix .\\apps\\web run build 通过；scripts/check.ps1 通过。",
    },
    {
        "date": "2026-06-23",
        "area": "在线阅读体验",
        "feature": "单页/连续滚动阅读模式",
        "details": "网页端阅读器新增 ReaderMode 通用状态，本地书库阅读器和远程在线阅读器共用“单页/连续”模式切换，并通过 localStorage 记住用户选择。连续模式使用统一的页面契约渲染可滚动页堆，当前页高亮，未加载页显示稳定占位；预取窗口在连续模式下自动扩大到当前页前 2 页、后 8 页，避免滚动时频繁空白。该能力放在阅读器层，不要求单个源站适配器实现定制 UI。",
        "status": "已完成",
        "verification": "npm --prefix .\\apps\\web run build 通过；node --check services/dev-api/server.mjs 与 services/dev-api/source-registry.mjs 通过；scripts/check.ps1 通过；Word 文档已重新生成并完成逐页渲染检查。",
    },
    {
        "date": "2026-06-24",
        "area": "在线阅读体验",
        "feature": "连续滚动当前页识别与进度同步",
        "details": "网页端本地书库阅读器和远程在线阅读器共用滚动视口同步逻辑：连续模式下每一页带有稳定 page 标记，阅读器滚动容器通过 IntersectionObserver 判断当前可见页，并同步页码输入框、进度条、缩略图高亮和当前页边框。进度保存采用 650ms 防抖，本地书库更新 shelf.last_page，远程会话调用 PATCH progress；该逻辑仍位于通用阅读器层，源站 adapter 不需要关心滚动 UI 或进度持久化。",
        "status": "已完成",
        "verification": "npm --prefix .\\apps\\web run build 通过；scripts/check.ps1 通过，覆盖 JSON 配置、源站注册表、Python bridge、Rust format、Rust workspace 与 Next.js 构建。",
    },
    {
        "date": "2026-06-24",
        "area": "在线阅读体验",
        "feature": "阅读器图片加载状态与失败重试",
        "details": "本地书库阅读器和远程在线阅读器的主图统一改为 renderReaderImage 渲染，单页模式和连续模式都能显示加载中占位、图片加载失败提示和“重试”按钮。连续模式的页容器从嵌套 button 调整为可聚焦的 role=button 容器，保留点击页图跳转能力，同时允许失败提示层内放置合法的重试按钮。重试会为当前图片请求追加刷新参数并重新加载，不修改页列表、阅读进度或源站 adapter 契约。",
        "status": "已完成",
        "verification": "npm --prefix .\\apps\\web run build 通过；scripts/check.ps1 通过，覆盖 JSON 配置、源站注册表、Python bridge、Rust format、Rust workspace 与 Next.js 构建。",
    },
    {
        "date": "2026-06-24",
        "area": "源站适配器",
        "feature": "新增 e-hentai.org 源站适配器",
        "details": "新增 scripts/ehentai_bridge.py，并在 config/source-adapters.json 注册 e-hentai 内置源站。该适配器复用 source_bridge_core 的 HTTP、cookie/header、图片校验、文件命名和 JSON 写入能力，支持 search、gallery、download、retry_folder、page_list、page_image 与 online_read；搜索解析 /g/{gid}/{token}/，画廊页解析命名空间 tag、页数和 /s/{page_token}/{gid}-{index} 页面，单页图片优先读取 img#img。E-Hentai 的 female:、language:、artist: 等命名空间 tag 会保留给 f_search，和 18comic 的普通关键词归一化策略分开。",
        "status": "已完成",
        "verification": "python -m py_compile scripts\\ehentai_bridge.py scripts\\source_bridge_core.py scripts\\check_source_adapters.py 通过；python scripts\\ehentai_bridge.py self-test 通过；python scripts\\check_source_adapters.py 返回 3 个 source adapter；node --check services/dev-api/source-registry.mjs 与 services/dev-api/server.mjs 通过。",
    },
    {
        "date": "2026-06-17",
        "area": "源站可用性与搜索体验",
        "feature": "默认多源搜索跳过当前不可用源站",
        "details": "排查发现前端黄条来自 18comic 在当前环境下公开搜索入口全部返回 403，而 Fangliding 正常返回结果。新增 source descriptor 的 available_for_default 与 unavailable_reason 字段，配置中声明 18comic 在未配置 COMIC18_COOKIE_FILE 或 COMIC18_HEADERS_FILE 时仍可手动选择，但不参与默认“全部可用源站”搜索。网页端默认来源从“全部启用源站”改为“全部可用源站”，下拉框对需要手动选择的源标记“手动”，并对历史合并任务中仅来自默认不可用源的黄条做降噪处理；手动指定单源失败时仍会显示明确错误。",
        "status": "已完成",
        "verification": "node --check services/dev-api/source-registry.mjs 与 services/dev-api/server.mjs 通过；python scripts/check_source_adapters.py 通过；/v1/sources 隔离烟测确认 Fangliding available_for_default=true、18comic available_for_default=false 且返回原因；npm --prefix .\\apps\\web run build 通过；scripts/check.ps1 通过。",
    },
    {
        "date": "2026-06-25",
        "area": "在线阅读体验",
        "feature": "直链源站自动识别与单页失败诊断",
        "details": "修复在线阅读器只显示坏图、错误原因不清楚的问题。开发期 API shim 的远程阅读会话创建支持在未传 source_id 时根据 gallery_url 域名自动识别支持 online_read 的源站；手动选择源站但 URL 域名明显不匹配时会返回明确错误。单页图片接口会记录每个 reader session/page 的失败原因，并新增 GET /v1/reader/sessions/{id}/pages/{index}/status 状态接口，返回 pending、ready 或 failed。网页端直链阅读不再逐个源站盲试，而是先匹配源站；图片加载失败时会查询状态接口，在阅读器失败层显示 HTTP 403、未解析到图片、非图片响应等具体诊断，并保留单页重试。",
        "status": "已完成",
        "verification": "node --check services/dev-api/server.mjs 通过；npm --prefix .\\apps\\web run build 通过；scripts/check.ps1 通过，覆盖 JSON 配置、源站注册表、Python bridge、Rust format、Rust workspace 与 Next.js 构建。",
    },
    {
        "date": "2026-06-25",
        "area": "在线阅读体验",
        "feature": "阅读器单页缓存强制刷新",
        "details": "修复远程在线阅读器重试按钮只刷新前端图片 URL、后端仍可能返回同一份坏缓存的问题。开发期 API shim 的单页图片接口现在识别 reader_retry 或 refresh=1 参数，重试时只删除当前页在 .data/page-cache 下的图片缓存，再重新调用源站 adapter 的 download-page；前端失败层的重试按钮继续复用统一图片状态组件，不改变页列表、阅读进度或源站适配器契约。",
        "status": "已完成",
        "verification": "node --check services/dev-api/server.mjs 通过；npm --prefix .\\apps\\web run build 通过；后续 scripts/check.ps1 统一验证。",
    },
    {
        "date": "2026-06-25",
        "area": "在线阅读体验",
        "feature": "远程阅读页状态面板与失败页处理",
        "details": "开发期 API shim 新增 GET /v1/reader/sessions/{id}/pages/status 批量页状态接口，可按 offset/limit 返回当前窗口内每页 pending、ready 或 failed 状态。网页端远程在线阅读器会把批量状态与浏览器图片加载状态合并，用于标记缩略图、临近页码按钮和连续滚动页；底部新增当前页状态、已就绪/失败/待加载统计、刷新状态、重试失败页和跳过失败页动作。该功能仍位于 Reader Service 与通用阅读器层，源站 adapter 不需要实现额外 UI 逻辑。",
        "status": "已完成",
        "verification": "node --check services/dev-api/server.mjs 通过；npm --prefix .\\apps\\web run build 通过；后续 scripts/check.ps1 统一验证。",
    },
    {
        "date": "2026-06-25",
        "area": "在线阅读维护",
        "feature": "远程阅读记录管理与缓存清理",
        "details": "开发期 API shim 新增 DELETE /v1/reader/sessions/{id} 用于删除远程阅读记录，并新增 POST /v1/reader/sessions/{id}/cache/clear 用于按当前页或整本清理 .data/page-cache 下的阅读缓存。网页端最近在线阅读面板支持标题、来源、URL 和 tag 筛选，支持展开更多历史记录，单条记录可清缓存或删除；远程阅读器底部新增清当前页与清本书缓存按钮，清理后会同步重置前端图片状态并触发重新加载。该能力保持在 Reader Service 和通用 UI 层，后续 App 端可复用同一维护接口。",
        "status": "已完成",
        "verification": "node --check services/dev-api/server.mjs 通过；npm --prefix .\\apps\\web run build 通过；后续 scripts/check.ps1 统一验证。",
    },
    {
        "date": "2026-06-25",
        "area": "在线阅读体验",
        "feature": "远程在线阅读书签",
        "details": "开发期 API shim 的远程阅读会话新增 bookmarks 字段，并持久化到 reader-sessions.json；新增 POST /v1/reader/sessions/{id}/bookmarks 用于添加或更新页书签，新增 DELETE /v1/reader/sessions/{id}/bookmarks/{page} 用于删除页书签。网页端远程阅读器底部新增当前页加书签/取消书签按钮和书签跳转条，支持直接跳到已标记页并单独删除书签；最近在线阅读卡片会显示书签数量。该能力作为 Reader Service 用户阅读元数据实现，源站 adapter 不需要关心。",
        "status": "已完成",
        "verification": "node --check services/dev-api/server.mjs 通过；npm --prefix .\\apps\\web run build 通过；后续 scripts/check.ps1 统一验证。",
    },
    {
        "date": "2026-06-25",
        "area": "下载任务稳定性",
        "feature": "整本下载实时进度、超时与孤儿任务恢复",
        "details": "修复图库直链下载任务长时间停在 queued/running、output 为 null 且没有可见进度的问题。18comic 与 e-hentai bridge 的 download-gallery 现在在每页处理后向 stderr 输出 __COMIC_PLATFORM_PROGRESS__ 结构化进度；开发期 API shim 捕获该进度并实时发布 task_progressed，错误日志会过滤这些进度行。图库下载 bridge 增加默认 30 分钟超时，running 但没有对应 bridge 子进程且超过 5 分钟未更新的任务会自动标记 failed，提示用户重跑。网页端在存在 queued/running/paused 任务时每 4 秒轮询一次 /v1/tasks，作为 SSE 丢事件时的兜底同步。E-Hentai 与 18comic 的下载目标现在按页序号保留，不再按重复 image_url 去重；过小图片响应会被识别为占位图或拦截结果并计入失败，避免 30x30 小图被当作整本下载成功。API 收到 stopped 或完全没有可用页的整本下载报告时会把任务标记为 failed，而不是把失败报告包装成 completed。",
        "status": "已完成",
        "verification": "node --check services/dev-api/server.mjs 通过；python -m py_compile scripts\\ehentai_bridge.py scripts\\18comic_bridge.py scripts\\source_bridge_core.py 通过；python scripts\\ehentai_bridge.py self-test 通过；npm --prefix .\\apps\\web run build 通过；scripts/check.ps1 通过。",
    },
    {
        "date": "2026-06-28",
        "area": "下载性能",
        "feature": "源站整本下载可控并发",
        "details": "18comic 与 e-hentai bridge 的整本下载从完全串行改为小窗口并发。下载前会先按 start_page、end_page、max_pages_per_run 裁剪目标页，避免只下载部分页时仍解析整本所有页面；随后使用可控线程池并发解析页面图片地址，并发保存图片文件。失败记录、download_state.json 和 __COMIC_PLATFORM_PROGRESS__ 仍由主线程统一更新，避免多线程抢写状态文件。新增 --download-concurrency 参数，默认并发 3，源站环境变量 COMIC18_DOWNLOAD_CONCURRENCY / EHENTAI_DOWNLOAD_CONCURRENCY 可单独调节，开发期 API shim 也可通过 DEV_API_GALLERY_DOWNLOAD_CONCURRENCY 统一覆盖，取值会限制在 1 到 8 之间。",
        "status": "已完成",
        "verification": "python -m py_compile scripts\\source_bridge_core.py scripts\\18comic_bridge.py scripts\\ehentai_bridge.py 通过；python scripts\\18comic_bridge.py self-test 与 python scripts\\ehentai_bridge.py self-test 通过；node --check services/dev-api/server.mjs、npm --prefix .\\apps\\web run build 与 scripts/check.ps1 均已通过。",
    },
    {
        "date": "2026-06-28",
        "area": "源站可用性",
        "feature": "18comic 项目内会话配置",
        "details": "为 18comic 增加项目内 Cookie/Header 配置能力。开发期 API shim 新增 /v1/source-auth/18comic，用于查看、保存和清除用户自己的授权会话文件，保存位置固定在 <project-root>\\.data\\source-auth；runSourceBridge 会自动把这些项目内文件作为 COMIC18_COOKIE_FILE / COMIC18_HEADERS_FILE 传给 18comic bridge，/v1/sources 会在读取前刷新状态，并要求 *_FILE 环境变量真实指向文件，避免假可用。网页端来源栏新增 18comic 会话配置区，可粘贴 Cookie 或整段请求头，保存后立即刷新来源状态并清空输入框。该功能只支持正常授权会话，不绕过登录、年龄门、验证码、封禁或限流。",
        "status": "已完成",
        "verification": "node --check services/dev-api/server.mjs 与 services/dev-api/source-registry.mjs 通过；python -m py_compile scripts\\18comic_bridge.py scripts\\source_bridge_core.py 通过；npm --prefix .\\apps\\web run build 通过；临时启动 dev-api 18080 端口烟测 /v1/source-auth/18comic，确认保存假 Header 后 18comic available_for_default=true，删除后恢复未配置。",
    },
    {
        "date": "2026-06-25",
        "area": "文件库质量诊断",
        "feature": "下载目录健康状态",
        "details": "开发期 API shim 的文件库扫描新增 health 诊断摘要，按统一结构返回正常、需处理或异常状态，并统计期望页数、已保存图片数、缺页数、失败页记录、download_state 中断状态和疑似占位小图数量。网页端文件库表格、卡片和详情侧栏会直接展示健康状态；详情中可看到具体问题、疑似小图样本和最近 download_state 更新时间。文件库 API 与网页端筛选区新增 health 过滤，可按正常、需处理、异常或异常/需处理快速筛选；顶部统计卡片也能一键进入异常和需处理目录。该能力属于 Library Service 通用诊断层，后续 App 端可复用同一字段。",
        "status": "已完成",
        "verification": "node --check services/dev-api/server.mjs 通过；npm --prefix .\\apps\\web run build 通过；后续 scripts/check.ps1 统一验证。",
    },
    {
        "date": "2026-06-28",
        "area": "源站可用性",
        "feature": "18comic 当前入口与挑战页诊断",
        "details": "18comic 搜索适配器改为优先请求浏览器中实际出现的 /meiman?f_search=... 入口，并保留 female:futanari 这类命名空间 tag 作为站内查询；旧 /search/photos?search_query=... 入口仅作为兼容兜底。适配器现在可解析 /album 与 /photo 两种结果链接，并从邻近 HTML 中提取真实缩略图。HTTP 公共层新增 Cloudflare/验证码/浏览器验证页识别，遇到 Just a moment 等挑战页时会快速失败并说明真实原因，不再误导为单纯缺 Cookie。18comic 默认请求头改为导航型浏览器请求，User-Agent 会优先读取本机 Edge 安装版本，curl_cffi 可用时默认使用更接近当前浏览器的 chrome146 指纹。",
        "status": "已完成",
        "verification": "python -m py_compile scripts\\source_bridge_core.py scripts\\18comic_bridge.py 通过；python scripts\\18comic_bridge.py self-test 通过；本机 Edge 版本检测返回 149 系 UA。受当前联网审批/额度限制，本轮未重新执行实时 18comic 网络探测。",
    },
    {
        "date": "2026-07-11",
        "area": "源站可用性",
        "feature": "18comic 移动 API 搜索与下载链路",
        "details": "将 18comic 适配器从 Cloudflare 保护的网页抓取主链路改为官方 JM 移动 API 优先，并保留原网页实现作为显式兼容兜底。新增项目本地 jmcomic 2.7.1 依赖和 .cache/python 加载路径；搜索、图集详情、章节图片列表、在线阅读单页缓存与整本下载全部映射回既有 source-neutral bridge 契约。API 页面使用 jmapi://photo/{photo_id}/{index} 内部描述符，下载时解析并完成图片解码。18comic 不再要求 Cookie/Header 才加入默认搜索，网页会话面板改为明确的可选备用配置。",
        "status": "已完成",
        "verification": "真实查询 chinese big breasts 返回 5 条 18comic 结果；首条图集解析为 68 页；bridge 与隔离 Node API 均成功下载首张 231104 字节 WebP；前端真实搜索任务状态为 completed；scripts/check.ps1 全部通过。",
    },
    {
        "date": "2026-07-12",
        "area": "在线阅读体验",
        "feature": "可收起底栏与沉浸式连续阅读",
        "details": "在线阅读器和本地阅读器的底部工具栏现可通过顶部按钮、底栏按钮、右下角恢复按钮或 H 快捷键随时收起与展开，并使用浏览器本地存储记住选择；首次使用默认进入收起状态。底栏展开时限制最大高度并允许内部滚动，远程页状态操作改为紧凑自动换行，避免工具区持续挤压漫画。连续阅读模式移除左右翻页占位，漫画列居中限宽，增加平滑滚动、底部安全空间和手机端响应式布局。",
        "status": "已完成",
        "verification": "Next.js 生产构建和 TypeScript 检查通过；scripts/check.ps1 的 JSON、源站配置、Python bridge、自测、Rust format/workspace 与 Web build 全部通过。",
    },
    {
        "date": "2026-07-12",
        "area": "本地启动",
        "feature": "强制重启进程竞态修复",
        "details": "修复 .\\scripts\\dev.ps1 -Fresh 依次停止 Web 与 API 时的竞态：Web 退出会触发旧启动器的 finally 块连带停止 API，导致脚本随后再次停止同一 API PID 时抛出 NoProcessFound 并中断。脚本现在把已自然退出的目标视为成功停止，并等待 3000/8080 端口真正释放后再启动新服务。日常启动命令仍为 .\\scripts\\dev.ps1，只有需要强制加载最新代码时才使用 -Fresh。",
        "status": "已完成",
        "verification": "先启动一组 Web/API 服务，再由第二次 -Fresh 替换；日志命中 API listener already stopped 竞态分支后继续启动成功，旧启动器退出，新 Web 与 API 分别在 3000/8080 返回 HTTP 200；普通启动同时验证可复用现有服务。",
    },
]

ENTRIES.append(
    {
        "date": "2026-07-13",
        "area": "词条系统",
        "feature": "全局禁用词条与中文映射自动补全",
        "details": "搜索词条输入改为浏览器式自动补全：输入中文、英文标准词条或别名时显示中文说明与英文标准词条，支持方向键、回车、Tab 和鼠标选择。新增持久化的全局禁用词条列表；创建或重跑搜索时自动把标准词条展开为中文与别名，开发 API 和 Rust 任务链在合并结果前过滤匹配漫画，无标签摘要会先读取图库详情补齐标签。词典由 EhTagTranslation/DatabaseReleases 的 db.text.json 转换生成，保留来源、版本、提交哈希和许可说明。",
        "status": "已完成",
        "verification": "词典检查覆盖 1363 条映射且 Next.js 生产构建通过；隔离开发 API 使用真实 18comic 搜索验证 female:big breasts/巨乳等价排除，5 条结果排除 2 条并保留 3 条；scripts/check.ps1 全量检查通过。",
    }
)


ENTRIES.append(
    {
        "date": "2026-07-19",
        "area": "开源与隐私",
        "feature": "公开仓库隐私边界与发布检查",
        "details": "新增公开仓库隐私检查脚本、GitHub Actions 和 pre-push 钩子；本地运行数据、会话、凭据、证书、文档渲染产物及研究数据均通过忽略规则保留在本机。示例配置改用占位值，Docker Compose 强制从本地环境读取密码，文档中的机器绝对路径改为通用占位符。公开分支提交身份使用 GitHub noreply 邮箱，原有历史只保存在本地私密 bundle 中。",
        "status": "已完成",
        "verification": "公开仓库检查覆盖跟踪文件名、文本与 DOCX 元数据、常见密钥格式、本机路径、大文件和提交邮箱；完整项目检查与公开快照检查通过后再发布。",
    }
)


def set_east_asian_font(run, font_name: str = "Microsoft YaHei UI") -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    set_east_asian_font(run)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def configure_styles(document: Document) -> None:
    styles = document.styles
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei UI"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei UI")

    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.size = Pt(22)
    styles["Title"].font.bold = True
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.bold = True
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.bold = True


def build_document() -> Document:
    document = Document()
    configure_styles(document)

    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("漫画平台项目功能变更记录")

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"当前版本记录日期：{date.today().isoformat()}    工作区：<project-root>")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(90, 98, 110)
    set_east_asian_font(run)

    document.add_paragraph(
        "本文档用于记录 newwork 新架构下每次新增、调整或验证的功能。后续每次继续开发时，会同步追加本文件。"
    )

    document.add_heading("记录摘要", level=1)
    summary = document.add_table(rows=1, cols=2)
    summary.alignment = WD_TABLE_ALIGNMENT.CENTER
    summary.autofit = False
    summary.style = "Table Grid"
    set_cell_text(summary.rows[0].cells[0], "项目", True)
    set_cell_text(summary.rows[0].cells[1], "内容", True)
    for cell in summary.rows[0].cells:
        shade_cell(cell, "D9EAF7")

    summary_rows = [
        ("当前主线", "网页端先行，未来 App 与网页端共用 Rust API 后端。"),
        ("主要技术", "Next.js / React / TypeScript / Rust / Axum / PostgreSQL。"),
        ("本地约束", "依赖、数据库、缓存优先放在 <project-root> 内，避免占用 C 盘。"),
        ("当前验证", "默认内存模式 scripts/check.ps1 通过；PostgreSQL 数据库本体已验证；开发期文件库检索、详情、预览、内置阅读器、搜索结果源站缩略图、阅读器页码跳转、缩略图条、单页/连续阅读模式、连续滚动当前页识别、阅读器图片失败重试、阅读器单页缓存强制刷新、远程阅读页状态面板、远程阅读记录管理、远程阅读书签、下载任务实时进度与超时保护、源站整本下载可控并发、18comic 项目内会话配置、下载目录健康诊断、开发启动脚本端口复用与强制重启、e-hentai 源站适配器、补缺入口、CBZ/PDF 导出、导出历史与网页下载均已完成验证。"),
    ]
    for key, value in summary_rows:
        row = summary.add_row().cells
        set_cell_text(row[0], key, True)
        set_cell_text(row[1], value)

    document.add_heading("功能变更明细", level=1)
    feature_table = document.add_table(rows=1, cols=6)
    feature_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    feature_table.autofit = False
    feature_table.style = "Table Grid"
    headers = ["日期", "模块", "功能", "说明", "状态", "验证"]
    for index, header in enumerate(headers):
        set_cell_text(feature_table.rows[0].cells[index], header, True)
        shade_cell(feature_table.rows[0].cells[index], "D9EAF7")

    for entry in ENTRIES:
        row = feature_table.add_row().cells
        set_cell_text(row[0], entry["date"])
        set_cell_text(row[1], entry["area"])
        set_cell_text(row[2], entry["feature"])
        set_cell_text(row[3], entry["details"])
        set_cell_text(row[4], entry["status"])
        set_cell_text(row[5], entry["verification"])

    widths = [0.8, 0.9, 1.35, 4.15, 0.65, 2.45]
    for row in feature_table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)

    paragraph = document.add_paragraph()
    paragraph.add_run("后续维护规则：").bold = True
    paragraph.add_run(" 每次新增功能、修改架构、接入基础设施或完成验证后，在本文件追加一条记录，并写清楚影响范围与验证方式。")

    return document


def main() -> None:
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
