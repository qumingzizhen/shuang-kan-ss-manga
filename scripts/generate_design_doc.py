from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
SOURCE = PROJECT_ROOT / "docs" / "项目架构与数据库设计.md"
OUTPUT = PROJECT_ROOT / "docs" / "项目架构与数据库设计.docx"


def apply_document_style(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)

    for style_name, size in [("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)]:
        style = document.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(74, 45, 92)


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = document.add_table(rows=1, cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    header = table.rows[0].cells
    for index, value in enumerate(rows[0]):
        paragraph = header[index].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(value)
        run.bold = True

    for row in rows[1:]:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value


def add_code_block(document: Document, lines: list[str]) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = document.styles["Normal"]
    for line in lines:
        run = paragraph.add_run(line + "\n")
        run.font.name = "Consolas"
        run.font.size = Pt(9)


def parse_markdown_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        raw = lines[index].strip().strip("|")
        cells = [cell.strip() for cell in raw.split("|")]
        if not all(set(cell) <= {"-", ":", " "} for cell in cells):
            rows.append(cells)
        index += 1
    return rows, index


def build_docx() -> None:
    document = Document()
    apply_document_style(document)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    index = 0
    in_code = False
    code_lines: list[str] = []

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(document, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue

        if in_code:
            code_lines.append(line)
            index += 1
            continue

        if not stripped:
            index += 1
            continue

        if stripped.startswith("|"):
            table_rows, next_index = parse_markdown_table(lines, index)
            add_table(document, table_rows)
            index = next_index
            continue

        if stripped.startswith("# "):
            document.add_heading(stripped[2:].strip(), level=0)
        elif stripped.startswith("## "):
            document.add_heading(stripped[3:].strip(), level=1)
        elif stripped.startswith("### "):
            document.add_heading(stripped[4:].strip(), level=2)
        elif stripped.startswith("- "):
            document.add_paragraph(stripped[2:].strip(), style="List Bullet")
        elif stripped.startswith(tuple(f"{number}. " for number in range(1, 10))):
            document.add_paragraph(stripped.split(". ", 1)[1], style="List Number")
        else:
            document.add_paragraph(stripped)

        index += 1

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_docx()
